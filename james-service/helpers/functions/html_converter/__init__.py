# import pdfkit
from weasyprint import HTML
import re
import io
from docx import Document

def add_element(element, text):
	return f"<{element}>{text}</{element}>"

def add_single_element(element):
	return f"<{element}>"


class HTMLConverter:

	regex = "<[<^]+?>"

	def __init__(self, json_project):
		self.json_project = json_project
		self.name = json_project['title']

	def get_docx(self):
		self.document = Document()
		self.document.add_heading(self.json_project['title']).add_run().add_break()

		for section in self.json_project['sections']:
			self.document.add_heading(section['title'])

			for paragraph in section['paragraphs']:
				paragraph = self.document.add_paragraph(paragraph['text'])
			paragraph.add_run().add_break()
			paragraph.add_run().add_break()

		target_stream = io.BytesIO()
		self.document.save(target_stream)
		target_stream.seek(0)
		return target_stream

	def get_html(self):
		self.html = add_element('h1', self.json_project['title'])
		self.html += add_single_element('hr') + add_single_element('br')

		for section in self.json_project['sections']:
			self.html += add_element('h2', section['title']) + add_single_element('br')
			for paragraph in section['paragraphs']:
				self.html += add_element('p', paragraph['text'])

			self.html += add_single_element('br') * 2

		return self.html

	def get_txt(self):
		self.text = self.name + "\n" + 24*"-" + "\n\n"

		for section in self.json_project['sections']:
			self.text += re.sub(self.regex, "" , section['title']) + "\n\n"
			for paragraph in section['paragraphs']:
				self.text += re.sub(self.regex, "" , paragraph['text'])
				self.text += "\n"

			self.text += "\n\n\n"

		return io.BytesIO(self.text.encode('utf-8'))
		# return self.text

	def get_pdf(self):
		pdf = HTML(string= self.get_html().encode('utf-8'))
		return io.BytesIO(pdf.write_pdf())

	def get(self, convert_type):
		return self.name, getattr(self, 'get_%s' % convert_type)()